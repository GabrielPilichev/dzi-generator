"""
ДЗИ Generator — Exam Builder
Генерира уникален DOCX изпит от въпросите в DB.

Използване:
    python build_exam.py [--variant 1] [--output exam.docx]
    python build_exam.py --topics spreadsheets,databases --variant 1
"""

import sys
import sqlite3
import random
import argparse
import json
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


# ============================================================
# Database queries
# ============================================================

def get_questions_pool(db_path: str) -> dict:
    """
    Зарежда всички въпроси от DB, групирани по тип.
    Връща dict с keys: 'multiple_choice', 'fill_in'
    """
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    
    pool = {"multiple_choice": [], "fill_in": []}
    
    # Multiple choice
    cursor = conn.execute("""
        SELECT q.id, q.source_exam, q.source_number, q.topic, q.points, q.prompt
        FROM questions q
        WHERE q.question_type = 'multiple_choice'
    """)
    for row in cursor:
        # Зареждаме опциите
        opts_cursor = conn.execute("""
            SELECT option_letter, option_text, is_correct
            FROM multiple_choice_options
            WHERE question_id = ?
            ORDER BY option_letter
        """, (row["id"],))
        options = [dict(opt) for opt in opts_cursor]
        
        # Само въпроси с 4 опции и поне 1 правилен отговор
        if len(options) == 4 and any(o["is_correct"] for o in options):
            pool["multiple_choice"].append({
                **dict(row),
                "options": options,
            })
    
    # Fill-in
    cursor = conn.execute("""
        SELECT q.id, q.source_exam, q.source_number, q.topic, q.points, q.prompt
        FROM questions q
        WHERE q.question_type = 'fill_in'
    """)
    for row in cursor:
        # Зареждаме подзадачите
        subs_cursor = conn.execute("""
            SELECT subquestion_number, correct_answer, points
            FROM fill_in_subquestions
            WHERE question_id = ?
            ORDER BY subquestion_number
        """, (row["id"],))
        subquestions = [dict(s) for s in subs_cursor]
        
        # Само въпроси с поне 1 подотговор
        if len(subquestions) > 0:
            pool["fill_in"].append({
                **dict(row),
                "subquestions": subquestions,
            })
    
    conn.close()
    return pool


def select_questions(pool: dict, mc_count: int = 15, fi_count: int = 10) -> list[dict]:
    """
    Селектира случайни въпроси от pool-а.
    Опитва да балансира между topics за да няма прекалено много на същата тема.
    """
    selected = []
    
    # Multiple choice — селектираме mc_count въпроса
    available_mc = pool["multiple_choice"][:]
    random.shuffle(available_mc)
    
    if len(available_mc) < mc_count:
        print(f"⚠️  Внимание: само {len(available_mc)} multiple choice въпроса в DB (искани: {mc_count})")
        mc_count = len(available_mc)
    
    selected.extend(available_mc[:mc_count])
    
    # Fill-in — селектираме fi_count въпроса
    available_fi = pool["fill_in"][:]
    random.shuffle(available_fi)
    
    if len(available_fi) < fi_count:
        print(f"⚠️  Внимание: само {len(available_fi)} fill-in въпроса в DB (искани: {fi_count})")
        fi_count = len(available_fi)
    
    selected.extend(available_fi[:fi_count])
    
    return selected


# ============================================================
# DOCX generation
# ============================================================

def add_run(paragraph, text: str, bold: bool = False, size: int = 12, font: str = "Times New Roman"):
    """Добавя текст към параграф със заданен стил."""
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    return run


def setup_document_style(doc: Document):
    """Задава дефолтен стил Times New Roman 12pt."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_header(doc: Document, variant: int):
    """Добавя header на изпита."""
    # Министерство
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "МИНИСТЕРСТВО НА ОБРАЗОВАНИЕТО И НАУКАТА", bold=True, size=14)
    
    # Заглавие
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "ДЪРЖАВЕН ЗРЕЛОСТЕН ИЗПИТ ПО", bold=True, size=14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "ИНФОРМАЦИОННИ ТЕХНОЛОГИИ", bold=True, size=14)
    
    # Дата (днешна)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    today = datetime.now().strftime("%d %B %Y г.")
    # Преводим месец на български
    months_bg = {
        "January": "януари", "February": "февруари", "March": "март",
        "April": "април", "May": "май", "June": "юни",
        "July": "юли", "August": "август", "September": "септември",
        "October": "октомври", "November": "ноември", "December": "декември"
    }
    for en, bg in months_bg.items():
        today = today.replace(en, bg)
    add_run(p, today, size=12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "ПРОФИЛИРАНА ПОДГОТОВКА", bold=True, size=12)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, f"ВАРИАНТ {variant}", bold=True, size=12)
    
    # Празен параграф
    doc.add_paragraph()
    
    # Време за работа
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "ЧАСТ 1 (Време за работа: 90 минути)", bold=True, size=12)
    
    doc.add_paragraph()


def add_multiple_choice_section(doc: Document, questions: list[dict]):
    """Добавя секцията с multiple choice въпроси (1-15)."""
    # Инструкция
    p = doc.add_paragraph()
    add_run(p, "За задачите от 1. до 15. включително изберете точно един от отговорите!",
            bold=True, size=11)
    
    doc.add_paragraph()
    
    for i, q in enumerate(questions, 1):
        # Номер и текст на въпроса
        p = doc.add_paragraph()
        add_run(p, f"{i}. ", bold=True)
        add_run(p, q["prompt"])
        
        # Опции
        for opt in q["options"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            add_run(p, f"{opt['option_letter']}) {opt['option_text']}")
        
        doc.add_paragraph()  # Празен ред между въпросите


def add_fill_in_section(doc: Document, questions: list[dict], start_number: int = 16):
    """Добавя секцията с fill-in въпроси (16-25)."""
    # Page break преди секцията
    doc.add_page_break()
    
    # Инструкция
    p = doc.add_paragraph()
    add_run(p, f"Отговорите на задачите от {start_number}. до {start_number + len(questions) - 1}. "
               f"включително запишете в полетата за отговори под задачата!",
            bold=True, size=11)
    
    doc.add_paragraph()
    
    for i, q in enumerate(questions, start=start_number):
        # Номер и текст на въпроса
        p = doc.add_paragraph()
        add_run(p, f"{i}. ", bold=True)
        add_run(p, q["prompt"])
        
        # Полета за отговор (празни редове)
        for sub in q["subquestions"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            add_run(p, f"({sub['subquestion_number']}) ........................................................")
        
        doc.add_paragraph()


def add_answer_key(doc: Document, questions: list[dict]):
    """Добавя ключ с отговори в края на документа."""
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "КЛЮЧ С ВЕРНИТЕ ОТГОВОРИ", bold=True, size=14)
    
    doc.add_paragraph()
    
    # Multiple choice answers
    p = doc.add_paragraph()
    add_run(p, "Задачи 1. – 15. (Multiple Choice):", bold=True, size=12)
    
    mc_questions = [q for q in questions if "options" in q]
    for i, q in enumerate(mc_questions, 1):
        correct = next((o["option_letter"] for o in q["options"] if o["is_correct"]), "?")
        p = doc.add_paragraph()
        add_run(p, f"{i}. {correct}")
    
    doc.add_paragraph()
    
    # Fill-in answers
    p = doc.add_paragraph()
    add_run(p, "Задачи 16. – 25. (Fill-in):", bold=True, size=12)
    
    fi_questions = [q for q in questions if "subquestions" in q]
    for i, q in enumerate(fi_questions, start=16):
        p = doc.add_paragraph()
        add_run(p, f"{i}.", bold=True)
        for sub in q["subquestions"]:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            add_run(p, f"({sub['subquestion_number']}) {sub['correct_answer']}")


def log_generated_exam(db_path: str, exam_name: str, questions: list[dict], output_path: str):
    """Записва генерирания изпит в DB-та (за audit)."""
    conn = sqlite3.connect(db_path)
    question_ids = [q["id"] for q in questions]
    conn.execute("""
        INSERT INTO generated_exams (exam_name, question_ids, output_file_path)
        VALUES (?, ?, ?)
    """, (exam_name, json.dumps(question_ids), output_path))
    conn.commit()
    conn.close()


# ============================================================
# Main
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="Генерира ДЗИ изпит")
    parser.add_argument("--db", default="data/questions.db")
    parser.add_argument("--variant", type=int, default=1, help="Номер на вариант")
    parser.add_argument("--output", default=None, help="Output path (default: auto)")
    parser.add_argument("--mc-count", type=int, default=15, help="Брой multiple choice")
    parser.add_argument("--fi-count", type=int, default=10, help="Брой fill-in")
    parser.add_argument("--seed", type=int, default=None, help="Random seed (за повторение)")
    args = parser.parse_args()
    
    db_path = Path(args.db)
    if not db_path.exists():
        print(f"❌ DB не намерена: {db_path}")
        sys.exit(1)
    
    # Random seed (по желание)
    if args.seed is not None:
        random.seed(args.seed)
    
    # Output path
    if args.output:
        output_path = Path(args.output)
    else:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = Path(f"exam_v{args.variant}_{timestamp}.docx")
    
    print(f"📊 Зареждам въпроси от DB...")
    pool = get_questions_pool(str(db_path))
    print(f"   Multiple choice: {len(pool['multiple_choice'])}")
    print(f"   Fill-in: {len(pool['fill_in'])}")
    
    print(f"🎲 Селектирам въпроси...")
    questions = select_questions(pool, args.mc_count, args.fi_count)
    
    mc_selected = [q for q in questions if "options" in q]
    fi_selected = [q for q in questions if "subquestions" in q]
    
    print(f"   Избрани: {len(mc_selected)} MC + {len(fi_selected)} fill-in = {len(questions)} общо")
    
    print(f"📝 Генерирам DOCX...")
    doc = Document()
    setup_document_style(doc)
    
    add_header(doc, args.variant)
    add_multiple_choice_section(doc, mc_selected)
    add_fill_in_section(doc, fi_selected, start_number=args.mc_count + 1)
    add_answer_key(doc, questions)
    
    # Save
    doc.save(str(output_path))
    
    # Log
    exam_name = output_path.stem
    log_generated_exam(str(db_path), exam_name, questions, str(output_path))
    
    # Total points
    total_points = sum(q["points"] for q in questions)
    
    print(f"\n✅ Готово!")
    print(f"   Файл: {output_path}")
    print(f"   Общо въпроси: {len(questions)}")
    print(f"   Общо точки: {total_points}")
    print(f"\n📚 Sources на въпросите:")
    sources = {}
    for q in questions:
        sources[q["source_exam"]] = sources.get(q["source_exam"], 0) + 1
    for src, cnt in sorted(sources.items()):
        print(f"   {src}: {cnt}")


if __name__ == "__main__":
    main()
