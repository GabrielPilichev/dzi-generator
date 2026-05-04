"""
Generate a topical worksheet (.docx) from approved questions.

Pulls questions from data/questions.db filtered by topic / area / class /
type / difficulty, formats them as a clean Times New Roman handout, and
optionally appends an answer key on a new page.

Production filter is always applied:
    WHERE is_ai_generated = 0 OR quality_score >= 1.0

Употреба:
    python3 src/build_worksheet.py --topic sumif --count 8
    python3 src/build_worksheet.py --topic sumif --topic vlookup --count 12
    python3 src/build_worksheet.py --area spreadsheets --class 11 --count 10
    python3 src/build_worksheet.py --topic sumif --types multiple_choice,true_false
    python3 src/build_worksheet.py --topic sumif --no-answers --seed 42
"""

from __future__ import annotations

import argparse
import random
import sqlite3
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt, Cm


FONT_NAME = "Times New Roman"
FONT_SIZE = 12
HEADER_SIZE = 16
SECTION_SIZE = 13


# ---------------------------------------------------------------------------
# DB layer
# ---------------------------------------------------------------------------

def fetch_questions(conn, *, topic_slugs, area_ids, class_num,
                    types, difficulty, count, seed):
    """Return up to `count` approved questions matching filters, randomised."""
    where = ["(q.is_ai_generated = 0 OR q.quality_score >= 1.0)"]
    params: list = []

    if topic_slugs:
        placeholders = ",".join("?" * len(topic_slugs))
        where.append(f"t.topic_slug IN ({placeholders})")
        params.extend(topic_slugs)

    if area_ids:
        placeholders = ",".join("?" * len(area_ids))
        where.append(f"a.area_id IN ({placeholders})")
        params.extend(area_ids)

    if class_num is not None:
        where.append(
            "EXISTS (SELECT 1 FROM topic_classes tc "
            "WHERE tc.topic_id = q.topic_id AND tc.class = ?)"
        )
        params.append(class_num)

    if types:
        placeholders = ",".join("?" * len(types))
        where.append(f"q.question_type IN ({placeholders})")
        params.extend(types)

    if difficulty is not None:
        where.append("q.difficulty = ?")
        params.append(difficulty)

    sql = f"""
        SELECT q.id, q.question_type, q.prompt, q.points, q.difficulty,
               q.has_image, q.image_path,
               t.title_bg AS topic_title, t.topic_slug
        FROM questions q
        LEFT JOIN curriculum_topics t ON t.id = q.topic_id
        LEFT JOIN curriculum_areas a   ON a.id = t.area_id
        WHERE {" AND ".join(where)}
    """

    rows = conn.execute(sql, params).fetchall()
    if not rows:
        return []

    rng = random.Random(seed)
    rng.shuffle(rows)
    return rows[:count]


def fetch_options(conn, question_id):
    return conn.execute(
        "SELECT option_letter, option_text, is_correct "
        "FROM multiple_choice_options "
        "WHERE question_id = ? "
        "ORDER BY option_letter",
        (question_id,),
    ).fetchall()


def fetch_subquestions(conn, question_id):
    return conn.execute(
        "SELECT subquestion_number, subquestion_text, "
        "       correct_answer, answer_alternatives, points "
        "FROM fill_in_subquestions "
        "WHERE question_id = ? "
        "ORDER BY subquestion_number",
        (question_id,),
    ).fetchall()


# ---------------------------------------------------------------------------
# Render layer
# ---------------------------------------------------------------------------

def _set_default_font(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE)


def _para(doc, text="", *, bold=False, size=None, align=None, after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if text:
        run = p.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(size or FONT_SIZE)
        run.bold = bold
    return p


def _heading(doc, text, *, size=HEADER_SIZE, after=8):
    return _para(doc, text, bold=True, size=size, after=after)


def render_header(doc, title_bg, count, points_total):
    _heading(doc, f"Работен лист — {title_bg}", size=HEADER_SIZE)

    info = doc.add_paragraph()
    info.paragraph_format.space_after = Pt(2)
    run = info.add_run(f"Дата: {date.today().strftime('%d.%m.%Y')}    "
                       f"Брой задачи: {count}    "
                       f"Общо точки: {points_total}")
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)

    name_line = doc.add_paragraph()
    name_line.paragraph_format.space_after = Pt(10)
    run = name_line.add_run("Име: ____________________________________    "
                            "Клас: ______    №: ______")
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)


def render_question(doc, conn, idx, q):
    qid, qtype, prompt, points, difficulty, has_image, image_path, topic_title, slug = q

    points = points or 1
    pts_label = f"({points} т.)" if points else ""
    prompt = (prompt or "").strip()

    # Question header line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{idx}. ")
    run.bold = True
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run = p.add_run(prompt)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    if pts_label:
        run = p.add_run(f"  {pts_label}")
        run.italic = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE - 1)

    if has_image and image_path:
        _para(doc, f"[изображение: {image_path}]", size=FONT_SIZE - 1)

    if qtype == "multiple_choice":
        for letter, text, _is_correct in fetch_options(conn, qid):
            _para(doc, f"   {letter}) {text}", after=2)

    elif qtype == "true_false":
        _para(doc, "   ☐ Вярно    ☐ Невярно", after=4)

    elif qtype == "fill_in":
        subs = fetch_subquestions(conn, qid)
        if subs:
            for num, sub_text, _ans, _alts, sub_pts in subs:
                line = f"   {num}) {sub_text}"
                if sub_pts:
                    line += f"   ({sub_pts} т.)"
                _para(doc, line, after=4)
        else:
            _para(doc, "   Отговор: " + "_" * 60, after=4)

    elif qtype == "short_answer":
        _para(doc, "   Отговор: " + "_" * 60, after=6)

    elif qtype in ("free_response", "practical"):
        for _ in range(4):
            _para(doc, "_" * 80, after=2)

    elif qtype == "matching":
        _para(doc, "   (свържете със стрелки)", after=6)

    else:
        _para(doc, "_" * 80, after=4)


def render_answer_key(doc, conn, items):
    doc.add_page_break()
    _heading(doc, "Ключ с отговори", size=SECTION_SIZE, after=8)

    for idx, q in enumerate(items, start=1):
        qid, qtype, prompt, *_ = q
        line = f"{idx}. "

        if qtype == "multiple_choice":
            opts = fetch_options(conn, qid)
            correct = [letter for letter, _t, is_c in opts if is_c]
            line += ", ".join(correct) if correct else "—"

        elif qtype == "fill_in":
            subs = fetch_subquestions(conn, qid)
            if subs:
                parts = []
                for num, _txt, ans, alts, _pts in subs:
                    val = ans or ""
                    if alts:
                        val = f"{val}  (или: {alts})" if val else alts
                    parts.append(f"{num}) {val or '—'}")
                line += "  |  ".join(parts)
            else:
                line += "—"

        elif qtype == "true_false":
            opts = fetch_options(conn, qid)
            correct = [letter for letter, _t, is_c in opts if is_c]
            line += correct[0] if correct else "(виж задачата)"

        else:
            line += "(отворен отговор — виж критерии)"

        _para(doc, line, after=2)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def parse_csv(s):
    return [x.strip() for x in s.split(",") if x.strip()] if s else []


def main() -> None:
    p = argparse.ArgumentParser(description="Generate a topical worksheet (.docx).")
    p.add_argument("--db", type=Path, default=Path("data/questions.db"))
    p.add_argument("--vault", type=Path, default=Path("vault"))
    p.add_argument("--topic", action="append", default=[],
                   help="topic_slug; repeatable")
    p.add_argument("--area", action="append", default=[],
                   help="area_id (e.g. spreadsheets); repeatable")
    p.add_argument("--class", dest="class_num", type=int, choices=range(8, 13))
    p.add_argument("--types", type=parse_csv, default=[],
                   help="comma-separated, e.g. multiple_choice,fill_in")
    p.add_argument("--difficulty", type=int, choices=range(1, 6))
    p.add_argument("--count", type=int, default=10)
    p.add_argument("--seed", type=int, default=None)
    p.add_argument("--title", default=None,
                   help="Override worksheet title")
    p.add_argument("--no-answers", action="store_true",
                   help="Skip answer key page")
    p.add_argument("--output", type=Path, default=None,
                   help="Output .docx path (default: data/worksheets/<slug>_<date>.docx)")
    p.add_argument("--dry-run", action="store_true")
    args = p.parse_args()

    if not args.db.exists():
        print(f"❌ DB not found: {args.db}")
        sys.exit(1)

    conn = sqlite3.connect(str(args.db))

    items = fetch_questions(
        conn,
        topic_slugs=args.topic,
        area_ids=args.area,
        class_num=args.class_num,
        types=args.types,
        difficulty=args.difficulty,
        count=args.count,
        seed=args.seed,
    )

    if not items:
        print("⚠️  No questions match these filters.")
        print("    Tried filters:")
        print(f"      topic={args.topic}  area={args.area}  class={args.class_num}")
        print(f"      types={args.types}  difficulty={args.difficulty}")
        sys.exit(2)

    print(f"📊 Fetched {len(items)} question(s).")

    # Title resolution
    if args.title:
        title_bg = args.title
    elif args.topic:
        # Use the first matching topic's pretty title
        row = conn.execute(
            "SELECT title_bg FROM curriculum_topics WHERE topic_slug = ?",
            (args.topic[0],),
        ).fetchone()
        title_bg = row[0] if row else ", ".join(args.topic)
    elif args.area:
        title_bg = ", ".join(args.area)
    else:
        title_bg = "смесен"

    points_total = sum((q[3] or 1) for q in items)

    if args.dry_run:
        print(f"📂 (dry-run) Would write worksheet: {title_bg}")
        for i, q in enumerate(items, 1):
            print(f"   {i:2d}. [{q[1]}] {(q[2] or '')[:70]}")
        conn.close()
        return

    # Build docx
    doc = Document()
    _set_default_font(doc)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    render_header(doc, title_bg, len(items), points_total)
    for i, q in enumerate(items, start=1):
        render_question(doc, conn, i, q)

    if not args.no_answers:
        render_answer_key(doc, conn, items)

    # Output path
    if args.output:
        out_path = args.output
    else:
        slug = (args.topic[0] if args.topic else
                args.area[0] if args.area else
                "mixed")
        out_dir = Path("data/worksheets")
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{slug}_{date.today().isoformat()}.docx"

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    conn.close()

    print(f"✅ Готово: {out_path}  ({len(items)} въпроса, {points_total} т.)")


if __name__ == "__main__":
    main()
